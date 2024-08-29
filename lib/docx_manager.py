import os
import docx
import docxcompose.composer
import lxml
from lxml.etree import Element
import lib.path as path
import lib.logs as logging
import subprocess
import pythoncom
import win32com.client

ABSOLUTE_FILE_PATH = os.path.abspath(__file__)

class DocxManager():

	################################################################# SURCHARGE

	@logging.log_call
	def __init__(self, **kwargs:dict) -> None:
		super().__init__()
		self._document = docx.Document()
		self._header_file = None
		self._encoding = "utf-8"
		self._export_path = path.Path()
		self._log_level = "info"
		self._pages = None
		self._path = path.Path()
		self._numbering = []
		self._saved_to_file = False

		self.path.update(ABSOLUTE_FILE_PATH.replace("py", "docx"))
		self.export_path.update(ABSOLUTE_FILE_PATH.replace("py", "pdf"))

	@logging.log_call
	def __str__(self) -> None:
		substrings = []
		substrings.append(f"{self.__name__}(")
		for attribute, value in vars(self).items():
			substrings.append(f"\t{attribute}: {str(value)}")
		substrings.append(f")")
		return "\n".join(substrings)

	################################################################### GETTERS

	@property
	def document(self) -> docx.Document:
		return self._document

	@property
	def encoding(self) -> str:
		return self._encoding

	@property
	def export_path(self) -> path.Path:
		return self._export_path

	@property
	def log_level(self) -> str:
		return self._log_level

	@property
	def numbering(self) -> list:
		return self._numbering

	@property
	def pages(self) -> list:
		return self._pages

	@property
	def path(self) -> path.Path:
		return self._path

	@property
	def saved_to_file(self) -> bool:
		return self._saved_to_file

	@property
	def header_file(self) -> str:
		return self._header_file
	
	################################################################### SETTERS

	@document.setter
	def document(self, document:docx.Document) -> None:
		self._document = document

	@encoding.setter
	def encoding(self, encoding:str) -> None:
		self._encoding = encoding

	@export_path.setter
	def export_path(self, export_path:str) -> None:
		self._export_path.update(export_path)

	@log_level.setter
	def log_level(self, log_level:str) -> None:
		self._log_level = log_level

	@pages.setter
	def pages(self, pages:list) -> None:
		self._pages = pages

	@numbering.setter
	def numbering(self, numbering:list) -> None:
		self._numbering = numbering

	@path.setter
	def path(self, path:str) -> None:
		self._path.update(path)

	@saved_to_file.setter
	def saved_to_file(self, saved_to_file:bool) -> None:
		self._saved_to_file = saved_to_file

	@header_file.setter
	def header_file(self, header_file:str) -> None:
		self._header_file = header_file
		if self.apply_style_from_header() == -1:
			self._header_file = None
	
	################################################################### METHODS

	@logging.log_call
	def append(self, path:str, heading_offset:int=None, anchor=None) -> bool:
		if self.save_to_file() == -1:
			logging.log(f'Impossible to append the content from the file at "{path}"', "Error")
			return -1		
		try:
			composer = docxcompose.composer.Composer(self.document)
			next_document = docx.Document(path)
			if heading_offset:
				for paragraph in next_document.paragraphs:
					if paragraph.style.name.startswith('Heading '):
						heading_level = int(paragraph.style.name.split(' ')[-1])
						new_heading_level = heading_level + heading_offset
						paragraph.style = f'Heading {new_heading_level}'
 		
	 		#Find the index of the anchor paragraph
			anchor_index = None
			if anchor:
				for i, paragraph in enumerate(composer.docx.paragraphs):
					if paragraph.text.strip() == anchor.strip():
						anchor_index = i
						break
			
			# Insert the content of the next document before the anchor
			if anchor_index is not None:
				composer.insert(new_document, index=anchor_index)
			else:
				# If anchor not found or not specified, append at the end
				composer.append(next_document)

			composer.save(self.path.abs)
			self.saved_to_file = False
		except Exception as e:
			logging.log(f'Error while appending the content from the file at "{path}" : {e}', "Error")
			return -1
		logging.log(f'Content from file at "{path}" concatenated to the document.')

	@logging.log_call
	def export(self, export_path:str=None) -> bool:
		export_path = export_path if export_path else self.export_path.abs
		if self.save_to_file() == -1:
			logging.log(f'Impossible to export the document at "{export_path}"', "Error")
			return -1
		try:
			word_app = win32com.client.gencache.EnsureDispatch("Word.Application")
			doc = word_app.Documents.Open(self.path.abs)
			# Documentation: https://learn.microsoft.com/en-us/office/vba/api/word.wdsaveformat
			doc.SaveAs2(export_path, FileFormat=17, ReadOnlyRecommended=True, EmbedTrueTypeFonts=True, SaveNativePictureFormat=True)
			doc.Close()
			word_app.Quit()
		except Exception as e:
			logging.log(f'Error while exporting the document at "{export_path}" : {e}', "Error")
			return -1
		logging.log(f'Document exported to "{self.export_path.abs}".')

	@logging.log_call
	def save_to_file(self, save_path:str=None) -> bool:
		if self.saved_to_file:
			logging.log(f'No changes to save on the document.', "Debug")
			return
		save_path = save_path if save_path else self.path.abs
		try:
			self.document.save(save_path)
			self.saved_to_file = True
		except Exception as e:
			logging.log(f'Error while saving the document at "{save_path}" : {e}', "Error")
			return -1
		logging.log(f'Document saved at "{save_path}".')

	@logging.log_call
	def break_page(self, anchor=None) -> bool:
		if anchor:
			paragraph = self.get_paragraph_with_text(anchor)
		else:
			paragraph = self.document.paragraphs[-1]
		try:
			new_run = paragraph.add_run()
			new_run.add_break(docx.enum.text.WD_BREAK.PAGE)
		except Exception as e:
			logging.log(f'Error while breaking last page of the document : {e}', "Error")
			return -1
		self.saved_to_file = False
		logging.log(f'Broke last page of the document.', "debug")

	@logging.log_call
	def add_page_number(self) -> bool:
		try:
			paragraph = self.document.sections[0].footer.paragraphs[0]
			run = paragraph.add_run()
			fldChar1 = docx.oxml.shared.OxmlElement('w:fldChar').set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
			run._r.append(fldChar1)
			instrText = docx.oxml.shared.OxmlElement('w:instrText').set(docx.oxml.shared.qn('xml:space'), 'preserve')
			instrText.text = "PAGE"
			run._r.append(instrText)
			fldChar2 = docx.oxml.shared.OxmlElement('w:fldChar').set(docx.oxml.shared.qn('w:fldCharType'), 'end')
			run._r.append(fldChar2)
			self.saved_to_file = False
		except Exception as e:
			logging.log(f'Error while adding page number : {e}', "Error")
			return -1
		logging.log(f'Page number added to the document.')

	@logging.log_call
	def load_file(self, path:str=None) -> bool:
		path = path if path else self.path.abs
		try:
			self.document = docx.Document(path)
			self.saved_to_file = False
		except Exception as e:
			logging.log(f'Error while loading the file at "{path}" : {e}', "Error")
			return -1
		logging.log(f'File at "{path}" loaded.')

	#@logging.log_call
	#def reload(self) -> bool:
	#	if (self.save_to_file() == -1) or (self.load_file() == -1):
	#		logging.log(f'Impossible to reload the document', "Error")
	#		return -1
	#	logging.log(f'Document reloaded.')

	@logging.log_call
	def clear_document(self) -> bool:
		try:
			for paragraph in self.document.paragraphs:
				p = paragraph._element
				p.getparent().remove(p)
				p._p = p._element = None
				self.saved_to_file = False
		except Exception as e:
			logging.log(f'Error while clearing the document : {e}', "Error")
			return -1
		logging.log(f'Document cleared.')

	@logging.log_call
	def apply_style_from_header(self) -> bool:
		if self.load_file(self.header_file) == -1:
			logging.log(f'Impossible to apply the style from the header at "{self.header_file}"', "Error")
			return -1
		logging.log(f'Style applied from the header at "{self.header_file}".') 

	@logging.log_call
	def update_table_of_contents(self) -> bool:
		self.save_to_file()
		try:
			# Create a Word application object
			word = win32com.client.gencache.EnsureDispatch("Word.Application")
			# Open the document
			doc = word.Documents.Open(self.path.abs)
			# Assuming there is only one TOC in the document (index 1)
			doc.TablesOfContents(1).Update()
			# Save the changes
			doc.Close(SaveChanges=True)
			# Exit the application
			word.Quit()
			self.saved_to_file = False
		except Exception as e:
			logging.log(f'Error while updating the table of contents : {e}', "Error")
			return -1
		# Reload the DOCX document with the changes
		self.load_file()
		logging.log(f'Table of contents updated.')

	@logging.log_call
	def update_table_of_illustrations(self) -> None:
		self.save_to_file()
		try:
			# Create a Word application object
			word = win32com.client.Dispatch("Word.Application")
			# Open the document
			doc = word.Documents.Open(self.path.abs)
			# Update the fields (captions, etc)
			doc.Fields.Update()
			# Assuming there is only one TOF in the document (index 1)
			doc.TablesOfFigures(1).Update()
			# Save the changes
			doc.Close(SaveChanges=True)
			# Exit the application
			word.Quit()
			self.saved_to_file = False
		except Exception as e:
			logging.log(f'Error while updating the table of figures : {e}', "Error")
			return -1
		# Reload the DOCX document with the changes
		self.load_file()
		logging.log(f'Table of figures updated.')

	@logging.log_call
	def open_export(self):
		subprocess.Popen(['start', '', self.export_path.abs], shell=True)
		logging.log(f'Exported document at "{self.export_path.abs}" opened.')  

	@logging.log_call
	def get_previous_paragraph(self, paragraph):
		#
		previous_paragraph = None
		#
		for current_paragraph in self.document.paragraphs:
			if current_paragraph.text == paragraph.text:
				return previous_paragraph
			previous_paragraph = current_paragraph

	@logging.log_call
	def get_paragraph_with_text(self, text):
		# Find the target paragraph
		for paragraph in self.document.paragraphs:
			if text in paragraph.text:
				return paragraph
		logging.log(f'Unable to find a paragraph containing the text "{text}".', "Warning")
		return -1

	@logging.log_call
	def insert_paragraph_after(self, previous_paragraph, new_paragraph):
		# Create a new paragraph element
		new_paragraph_element = new_paragraph._element
		
		# Remove the new paragraph element from its current position if it exists
		if new_paragraph_element.getparent() is not None:
			new_paragraph_element.getparent().remove(new_paragraph_element)
		
		# Insert the new paragraph element after the previous paragraph element
		previous_paragraph._element.addnext(new_paragraph_element)

	@logging.log_call
	def paragraph(self, text, style="Normal"):
		#
		return docx.Document().add_paragraph(text, style=style)

	@logging.log_call
	def text(self, text, style="Normal"):
		#
		return self.paragraph(text=text, style=style)

	@logging.log_call
	def link(self, text, url, style="Hyperlink"):
		paragraph = self.paragraph(text="", style=style)

		# This gets access to the document.xml.rels file and gets a new relation id value
		part = paragraph.part
		r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

		# Create the w:hyperlink tag and add needed values
		hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
		hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

		# Create a new run object (a wrapper over a 'w:r' element)
		new_run = docx.text.run.Run(
			docx.oxml.shared.OxmlElement('w:r'), paragraph)
		new_run.text = text
		#new_run.style = style

		# Join all the xml elements together
		hyperlink.append(new_run._element)
		paragraph._p.append(hyperlink)

		return paragraph

	@logging.log_call
	def add_paragraph(self, text="", style="Normal", anchor=None):
		#
		paragraph = self.document.add_paragraph(text, style=style)
		#
		if anchor:
			#
			anchor_paragraph = self.get_paragraph_with_text(anchor)
			#
			previous_paragraph = self.get_previous_paragraph(anchor_paragraph)
			#
			self.insert_paragraph_after(previous_paragraph, paragraph)
		#
		return paragraph

	@logging.log_call
	def title(self, text, level=1, anchor=None):
		#
		paragraph = self.add_paragraph(text, style=f"Heading {level}", anchor=anchor)
		#
		logging.log(f'Title {level} "{text}" added to the document.')

	@logging.log_call
	def add_text(self, text, style="Normal", anchor=None):
		#
		paragraph = self.add_paragraph(text, style=style, anchor=anchor)
		#
		logging.log(f'Text "{text}" added to the document.')

	@logging.log_call
	def add_link(self, text, url, style="Hyperlink", anchor=None):
		paragraph = self.add_paragraph(style=style, anchor=anchor)

		# This gets access to the document.xml.rels file and gets a new relation id value
		part = paragraph.part
		r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

		# Create the w:hyperlink tag and add needed values
		hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
		hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

		# Create a new run object (a wrapper over a 'w:r' element)
		new_run = docx.text.run.Run(
			docx.oxml.shared.OxmlElement('w:r'), paragraph)
		new_run.text = text
		#new_run.style = style

		# Join all the xml elements together
		hyperlink.append(new_run._element)
		paragraph._p.append(hyperlink)

		logging.log(f'Link to "{text}" added to the document.')

	@logging.log_call
	def bookmark(self, name, anchor=None):
		el = [el for el in self.document._element[0] if el.tag.endswith('}p')][-1]
		el.append(lxml.etree.Element(docx.oxml.shared.qn('w:bookmarkStart'),{docx.oxml.shared.qn('w:id'):'0',docx.oxml.shared.qn('w:name'):name}))
		el.append(lxml.etree.Element(docx.oxml.shared.qn('w:bookmarkEnd'),{docx.oxml.shared.qn('w:id'):'0'}))

	@logging.log_call
	def add_image(self, path, width=18.5, caption=None, alignment="center", anchor=None) -> bool:
		try:
			# Add picture to document
			paragraph = self.add_paragraph(anchor=anchor)
			#
			if anchor:
				#
				anchor_paragraph = self.get_paragraph_with_text(anchor)
				#
				previous_paragraph = self.get_previous_paragraph(anchor_paragraph)
				#
				self.insert_paragraph_after(previous_paragraph, paragraph)
			#
			run = paragraph.add_run()
			run.add_picture(path, width=docx.shared.Cm(width))

			# Set alignment
			if alignment == 'left':
				paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
			elif alignment == 'center':
				paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
			elif alignment == 'right':
				paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

			if caption:
				paragraph = self.add_paragraph(f'Illustration ', style='Caption', anchor=anchor)

				# numbering field
				run = paragraph.add_run()

				fldChar = docx.oxml.shared.OxmlElement('w:fldChar')
				fldChar.set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
				run._r.append(fldChar)

				instrText = docx.oxml.shared.OxmlElement('w:instrText')
				instrText.text = f' SEQ Illustration \\* ARABIC'
				run._r.append(instrText)

				fldChar = docx.oxml.shared.OxmlElement('w:fldChar')
				fldChar.set(docx.oxml.shared.qn('w:fldCharType'), 'end')
				run._r.append(fldChar)

				# caption text
				paragraph.add_run(f': {caption}')
		except Exception as e:
			logging.log(f'Error while exporting the image at "{self.export_path}" : {e}', "Error")
			return -1
		logging.log(f'Image at "{path}" concatenated to the document.')

	@logging.log_call
	def replace_text_in_paragraph(self, paragraph, old_text, new_text):
		if old_text in paragraph.text:
			paragraph.text = paragraph.text.replace(old_text, new_text)
			logging.log(f'Text "{old_text}" replaced with "{new_text}".')

	@logging.log_call
	def replace_text_in_table(self, table, old_text, new_text):
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					self.replace_text_in_paragraph(paragraph, old_text, new_text)

	@logging.log_call
	def replace_text(self, old_text, new_text):
		#	
		# Iterate through all the paragraphs in the document
		#
		for paragraph in self.document.paragraphs:
			self.replace_text_in_paragraph(paragraph, old_text, new_text)
		#	
		# Iterate through all the tables in the document (if any)
		#	
		for table in self.document.tables:
			self.replace_text_in_table(table, old_text, new_text)
		#	
		# Replace text in headers and footers
		#	
		for section in self.document.sections:
			#	
			for paragraph in section.header.paragraphs:
				#	
				self.replace_text_in_paragraph(paragraph, old_text, new_text)
			#	
			for paragraph in section.footer.paragraphs:
				#	
				self.replace_text_in_paragraph(paragraph, old_text, new_text)
		#	
		# Save the updates
		#	
		self.save_to_file()

		pythoncom.CoInitialize()
		# Open the Word application
		word = win32com.client.gencache.EnsureDispatch('Word.Application')
		doc = word.Documents.Open(self.path.abs)

		# Replace text in text boxes (shapes)
		for shape in doc.Shapes:
			if shape.TextFrame.HasText:
				text_range = shape.TextFrame.TextRange
				text_range.Find.Execute(old_text, False, False, False, False, False, True, 1, False, new_text, 2)

		# Save and close the document
		doc.SaveAs(self.path.abs)
		doc.Close(False)
		word.Application.Quit()

		# Reload the DOCX document with the changes
		self.load_file()

	#@logging.log_call
	#def table(self, data, border_color="#000000"):
	#	# Create a new Document
	#	doc = self.document
	#	
	#	# Add a table with the given data
	#	table = doc.add_table(rows=len(data), cols=len(data[0]))
	#
	#	# Define the RGB color from the hex code
	#	border_color_rgb = docx.shared.RGBColor(int(border_color[1:3], 16), int(border_color[3:5], 16), int(border_color[5:7], 16))
	#
	#	#for i, row_data in enumerate(data):
	#		#row = table.rows[i]
	#		#for j, cell_data in enumerate(row_data):
	#			#cell = row.cells[j]
	#			#cell.paragraphs = data[i][j][1]
	#			#cell.text = str(cell_data)
	#			# Center the text vertically
	#			#cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
	#			# Center the text horizontally
	#			#for paragraph in cell.paragraphs:
	#				#for run in paragraph.runs:
	#					#run.font.size = docx.shared.Pt(11)
	#				#paragraph.alignment = 1  # Center alignment
	#
	#	# Apply the border style to the table
	#	tbl = table._tbl  # Get the table element
	#	tbl_pr = tbl.tblPr  # Access table properties
	#	tbl_borders = docx.oxml.OxmlElement('w:tblBorders')  # Create a new borders element
	#
	#	# Define border attributes
	#	borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
	#	for border in borders:
	#		border_element = docx.oxml.OxmlElement(f'w:{border}')
	#		border_element.set(docx.oxml.ns.qn('w:val'), 'single')
	#		border_element.set(docx.oxml.ns.qn('w:sz'), '4')
	#		border_element.set(docx.oxml.ns.qn('w:space'), '0')
	#		border_element.set(docx.oxml.ns.qn('w:color'), border_color[1:])  # Use hex color without '#'
	#		tbl_borders.append(border_element)
	#
	#	tbl_pr.append(tbl_borders)  # Append borders to table properties
	#	
	#	#
	#	logging.log(f'Table {data} added to the document.')
	#
	#	#
	#	return table

	#@logging.log_call
	#def increase_numbering(self, increased_heading_level:int=1) -> None:
	#	increased_heading_level = int(increased_heading_level)
	#	while len(self.numbering) < increased_heading_level:
	#		self.numbering.append(0)
	#	self.numbering[increased_heading_level -1] += 1
	#	if len(self.numbering) > increased_heading_level:
	#		#
	#		#	Remove the useless "0"
	#		#	After an increase, [1, 1, 1] becomes [2] instead of [2, 0, 0]
	#		#
	#		self.numbering = self.numbering[:increased_heading_level]